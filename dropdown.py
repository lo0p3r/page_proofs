# coding: utf8

import kivy
kivy.require("1.11.0")

from kivy.app import App
from kivy.uix.gridlayout import GridLayout
from kivy.core.window import Window
from kivy.properties import StringProperty

from docx import Document
from docx.shared import Mm
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

import os
import glob

Window.clearcolor = (0.5, 0.5, 0.5, 1)


class CalcGridLayout(GridLayout):
    filePath = StringProperty('')

    def __init__(self, **kwargs):
        super(CalcGridLayout, self).__init__(**kwargs)
        Window.bind(on_dropfile=self._on_file_drop)

    def reduced_image(self):
        print(self.filePath)

    def _on_file_drop(self, window, file_path):
        print(file_path)
        self.filePath = file_path.decode("utf-8")     # convert byte to string
        self.ids.img.source = self.filePath
        self.ids.img.reload()   # reload image
    def build_word(self, image_path):
        doc_path = os.path.join(
            os.path.dirname(image_path),
            os.path.basename(image_path) + ".docx")
        print(doc_path)
        imagefiles = []
        for file in glob.glob(os.path.join(image_path,"*.jpg")):
            imagefiles.append(file)
        print(len(imagefiles))

        top_margin = Mm(5.1)
        left_margin = Mm(6.1)
        page_height = Mm(210)
        page_width = Mm(297)
        header_distance = Mm(0)
        footer_distance = Mm(0)
        image_height = Mm(210 - 2 * 5.1)
        image_width = Mm(148.5 - 6.1)

        document = Document()
        section = document.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_height = page_height
        section.page_width = page_width
        section.top_margin = top_margin
        section.bottom_margin = top_margin
        section.left_margin = left_margin
        section.right_margin = left_margin
        section.header_distance = header_distance
        section.footer_distance = footer_distance

        new_paragraph = True

        for image in imagefiles:
            if new_paragraph:
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            new_paragraph = not new_paragraph
            run = paragraph.add_run()
            run.add_picture(image, height=image_height) #, width=image_width)

        # sections = document.sections
        # for section in sections:
        #     section.orientation = WD_ORIENT.LANDSCAPE
        #     section.top_margin = top_margin
        #     section.bottom_margin = top_margin
        #     section.left_margin = left_magrin
        #     section.right_margin = left_magrin

        document.save(doc_path)
        print("done")

class DragDropApp(App):

    def build(self):
        return CalcGridLayout()


if __name__ == "__main__":
    DragDropApp().run()